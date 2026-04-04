#!/usr/bin/env python3
"""
md_to_docx: Markdown ファイルを Word (.docx) に変換する

対応要素:
  - frontmatter（--- で囲まれたブロック）→ 除外
  - # / ## / ### 見出し → Word 見出しスタイル (Heading 1/2/3)
  - **太字** → 太字ラン
  - --- 水平線 → 段落区切り（空行）
  - 通常テキスト → Normal スタイル

Usage:
  python3 md_to_docx.py --input [input.md] --output [output.docx]
  python3 md_to_docx.py --input [input.md]        # 同フォルダに .docx を生成
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("python-docx が見つかりません。pip install python-docx でインストールしてください。")


def strip_frontmatter(text: str) -> str:
    """--- で囲まれた先頭の frontmatter を除去する"""
    stripped = text.lstrip()
    if stripped.startswith("---"):
        end = stripped.find("\n---", 3)
        if end != -1:
            return stripped[end + 4:].lstrip("\n")
    return text


def parse_inline(text: str) -> list[tuple[str, bool]]:
    """
    インライン要素をパースして (テキスト, is_bold) のリストを返す。
    **太字** のみ対応。
    """
    parts: list[tuple[str, bool]] = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts


def add_paragraph_with_inline(doc: Document, text: str, style: str = "Normal") -> None:
    """インライン要素（太字）を含む段落を追加する"""
    p = doc.add_paragraph(style=style)
    for segment, is_bold in parse_inline(text):
        run = p.add_run(segment)
        if is_bold:
            run.bold = True


def convert(input_path: Path, output_path: Path) -> None:
    text = input_path.read_text(encoding="utf-8")
    text = strip_frontmatter(text)

    doc = Document()

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游明朝"
    font.size = Pt(10.5)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        # 空行
        if not line.strip():
            continue

        # 水平線 --- → 空の区切り段落
        if re.fullmatch(r"-{3,}", line.strip()):
            doc.add_paragraph()
            continue

        # 見出し
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            # ** を取り除いてからテキスト設定
            heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
            doc.add_heading(heading_text, level=level)
            continue

        # 通常行（インライン太字含む）
        add_paragraph_with_inline(doc, line)

    doc.save(str(output_path))
    print(f"✅ 変換完了: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Markdown → Word (.docx) 変換")
    parser.add_argument("--input", "-i", type=Path, required=True, help="入力 .md ファイルのパス")
    parser.add_argument("--output", "-o", type=Path, required=False, help="出力 .docx ファイルのパス（省略時: 入力と同フォルダ）")
    args = parser.parse_args()

    input_path = args.input.resolve()
    if not input_path.exists():
        print(f"❌ ファイルが見つかりません: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output.resolve() if args.output else input_path.with_suffix(".docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"変換中: {input_path.name} → {output_path.name}")
    convert(input_path, output_path)


if __name__ == "__main__":
    main()
